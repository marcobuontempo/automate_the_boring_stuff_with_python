#! python3
# custom_invitations.py - Generates a Word document with custom invitations, using the names from a list contained in 'guests.txt'

import sys
import docx

# populate list of guest names
guest_names = []
try:
    guest_file = open('guests.txt')
    guest_names = guest_file.readlines()
except:
    # exit if error loading names
    print("Error: 'guests.txt' does not exist in current directory or was unable to load")
    sys.exit()

# load styles document
invitation_document = docx.Document("styles.docx")
for i,guest_name in enumerate(guest_names):
    guest_name = guest_name.strip() # remove \n character from each guest name
    # write contents to document
    invitation_document.add_paragraph("It would be a pleasure to have the company of", style="Heading2")
    invitation_document.add_paragraph(guest_name, style="Heading1")
    invitation_document.add_paragraph("at 11010 Memory Lane on the Evening of", style="Heading2")
    invitation_document.add_paragraph("April 1st", style="Heading3")
    invitation_document.add_paragraph("at 7 o'clock", style="Heading2")
    # add page break after each invitation (except for last invitation)
    if i < len(guest_names)-1:
        invitation_document.add_page_break()

# save file and display success message
invitation_document.save("invites.docx")
print("Creation successful. Check 'invites.docx'.")